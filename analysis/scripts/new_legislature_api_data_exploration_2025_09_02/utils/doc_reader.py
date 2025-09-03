import argparse
import os
import re
import sys
from abc import ABC, abstractmethod
from functools import cached_property
from tempfile import TemporaryDirectory

import docx
from PyPDF2 import PdfReader as PyPDF2Reader

OCR_TEXT_HEADER_PATTERN = r"\n{0,2}---- OCR TEXT \d+\s*----\n?"


class DocumentReader(ABC):
    def __init__(self, filepath: str): ...
    @cached_property
    @abstractmethod
    def text(self) -> str:
        pass

    @cached_property
    @abstractmethod
    def num_pages(self) -> int:
        pass

    @cached_property
    @abstractmethod
    def file_text(self) -> str:
        pass

    @cached_property
    @abstractmethod
    def ocr_text(self) -> str:
        raise NotImplementedError


class PdfReader(DocumentReader):
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.reader = PyPDF2Reader(filepath)

    @cached_property
    def ocr_text_pages(self) -> list[str]:
        print(f"Extracting images for {self.filepath}")
        import pytesseract
        from pdf2image import (
            convert_from_path,  # pyright: ignore[reportUnknownVariableType]
        )

        with TemporaryDirectory() as tmpdir:
            ocr_text_pages: list[str] = [
                pytesseract.image_to_string(image)  # pyright: ignore[reportUnknownMemberType,reportAttributeAccessIssue]
                for image in convert_from_path(self.filepath, output_folder=tmpdir)
            ]
        return ocr_text_pages

    @cached_property
    def ocr_text(self) -> str:
        ocr_text = ""
        for i, page in enumerate(self.ocr_text_pages):
            if not page:
                continue
            if i:
                ocr_text += "\n"
            ocr_text += f"---- OCR TEXT {i}----\n{page}\n"
        return ocr_text

    @cached_property
    def file_text(self) -> str:
        return "\n".join(
            page.extract_text() or "" for page in self.reader.pages
        ).strip()

    @cached_property
    def text(self) -> str:
        return f"""---- PDF TEXT ----\n{self.file_text}\n\n{self.ocr_text}""".strip()

    @cached_property
    def num_pages(self) -> int:
        return len(self.reader.pages)


class DocxReader(DocumentReader):
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.ext = os.path.splitext(filepath)[1].lower()

    @cached_property
    def num_pages(self) -> int:
        if self.ext == ".docx":
            doc = docx.Document(self.filepath)
            return sum(p.contains_page_break for p in doc.paragraphs)
        raise NotImplementedError

    @cached_property
    def text(self) -> str:
        if self.ext == ".docx":
            doc = docx.Document(self.filepath)
            return "\n".join(p.text for p in doc.paragraphs)
        raise ValueError(f"Unsupported file type or missing parser: {self.ext}")

    @cached_property
    def file_text(self) -> str:
        return self.text

    @cached_property
    def ocr_text(self) -> str:
        return ""


class TxtReader(DocumentReader):
    def __init__(self, filepath: str):
        assert filepath.endswith(".txt")
        self.filepath = filepath
        self.ext = os.path.splitext(filepath)[1].lower()

    @cached_property
    def num_pages(self) -> int:
        raise NotImplementedError

    @cached_property
    def text(self) -> str:
        with open(self.filepath, "r") as f:
            return f.read()

    @cached_property
    def file_text(self) -> str:
        text = self.text
        if match := re.search(OCR_TEXT_HEADER_PATTERN, text):
            text = text[: match.span()[0]].strip()
            if match := re.search(r"---- PDF TEXT ----\n?", text):
                text = text[match.span()[1] :].strip()
        return text

    @cached_property
    def ocr_text(self) -> str:
        text = self.text
        if match := re.search(OCR_TEXT_HEADER_PATTERN, text):
            text = text[match.span()[1] :].strip()
        return re.sub(OCR_TEXT_HEADER_PATTERN, "", text).strip()


DOC_READERS: dict[str, type[DocumentReader]] = {
    ".pdf": PdfReader,
    ".docx": DocxReader,
    ".txt": TxtReader,
}


class DocumentRef:
    def __init__(self, filepath: str, *, use_text_backup: bool):
        ext = os.path.splitext(filepath)[1].lower()
        txt_filepath = filepath.removesuffix(ext) + ".txt"
        if (
            use_text_backup
            and not filepath.endswith(".txt")
            and os.path.exists(txt_filepath)
        ):
            self.filepath = txt_filepath
            self.ext = ".txt"
        else:
            self.ext = ext
            self.filepath = filepath
        if self.ext not in DOC_READERS:
            raise ValueError(f"Unsupported file type: {self.ext}")

    @cached_property
    def document(self) -> DocumentReader:
        return DOC_READERS[self.ext](self.filepath)

    def save_text_backup(self, overwrite: bool = False):
        if not overwrite and os.path.exists(self.txt_filepath):
            return
        with open(self.txt_filepath, "w") as f:
            f.write(self.text)

    @cached_property
    def txt_filepath(self) -> str:
        return self.filepath.removesuffix(self.ext) + ".txt"

    @cached_property
    def text(self) -> str:
        return self.document.text.strip()

    @cached_property
    def file_text(self) -> str:
        return self.document.file_text.strip()

    @cached_property
    def ocr_text(self) -> str:
        return self.document.ocr_text.strip()

    @cached_property
    def num_pages(self) -> int:
        return self.document.num_pages


def main():
    parser = argparse.ArgumentParser(
        description="Extract text from a PDF or DOC/DOCX file and output to stdout."
    )
    parser.add_argument("file", help="Path to the file")
    parser.add_argument("--no-ocr", action="store_true", help="Disable OCR")
    args = parser.parse_args()

    try:
        doc = DocumentRef(args.file, use_text_backup=False)
        if args.no_ocr:
            text = doc.file_text
        else:
            text = doc.text

        print(text)
    except Exception as e:
        print(f"Error reading file: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
